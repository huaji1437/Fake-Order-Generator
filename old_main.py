import re
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


学院名: str
标题格式化文本: str = "%s"

def set_cell_shading(cell, shade):
    """设置单元格底纹颜色"""
    tcPr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), shade)
    tcPr.append(shading)
def parse_student_data(input_str):
    """使用正则表达式解析学生数据"""
    pattern = r'(?:^\d+\.\s*)?(\d{2})(云计算|计算机应用技术|计应|大数据技术|大数据|网络技术|网络|软件技术|软件|人工智能技术应用|人工智能|数字媒体技术班|数字媒体技术|数字媒体|数媒|电竞)(五年制)?(单)?(?:(\d+)(?:班|班级)?)?\s*?(\S+)'
    matches = re.findall(pattern, input_str)
    students = []
    for match in matches:
        grade = match[0]  # 年级
        base_class = match[1]  # 基础专业名称
        is_five_year = match[2] if match[2] else ''  # 五年制标记
        is_single_class = match[3] if match[3] else ''  # 单班标记
        class_num = match[4] if match[4] else ''  # 阿拉伯数字班级号
        name = match[5].strip()  # 姓名
        modifiers = is_five_year + is_single_class
        class_mapping = {
            '网络技术': '网络',
            '计算机应用': '计应',
            '软件技术': '软件',
            '云计算': '云计算',
            '电子竞技': '电竞',
            '人工智能': '人工智能',
            '人工智能技术应用': '人工智能',
            '大数据技术': '大数据',
            '数字媒体': '数媒',
        }
        full_base_class = base_class
        for short, full in class_mapping.items():
            if short in base_class:
                full_base_class = full
                break
        class_name = full_base_class + modifiers + class_num
        full_class = f"{grade}{class_name}"
        students.append((full_class, name))
        print(f"{grade}{class_name} {name}")
    return students
def get_student_input():
    """获取用户输入的学生数据"""
    print("请输入学生数据（每行一个学生，格式如：23计应2xxx，输入空行结束）：")
    lines = []
    while True:
        line = input().strip()
        if not line:
            break
        lines.append(line)
    return "\n".join(lines)

def create_leave_form():
    global year, month, day, cause
    input_data = get_student_input()
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = '等线'
    font.size = Pt(11)

    font.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("浙江长征职业技术学院计信学院请假单")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = '等线'
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

    # 时间
    time_para = doc.add_paragraph()
    time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    time_run = time_para.add_run(f"□早自习 ☑晚自习 请假时间：{year}年{month}月{day}日")
    time_run.font.name = '楷体'
    time_run.font.size = Pt(10)
    time_run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

    greeting = doc.add_paragraph()
    greeting_run = greeting.add_run("各班级：")
    greeting_run.font.name = '仿宋'
    greeting_run.font.size = Pt(12)
    greeting_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    reason = doc.add_paragraph()
    reason.paragraph_format.first_line_indent = Inches(0.3)
    reason_run = reason.add_run(f"因{cause}工作需要，以下同学需请假。")
    reason_run.font.name = '仿宋'
    reason_run.font.size = Pt(12)
    reason_run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    students = parse_student_data(input_data)
    students.sort(key=lambda x: (x[0][:2], x[0][2:]))

    num_students = len(students)
    num_rows = (num_students + 1) // 2 + 1

    table = doc.add_table(rows=num_rows, cols=6)
    table.style = 'Table Grid'

    headers = ["序号", "班级", "姓名", "序号", "班级", "姓名"]
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = '等线'
                run.font.size = Pt(11)
                run.bold = True
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
        set_cell_shading(hdr_cells[i], "D9D9D9")

    for i in range(num_rows - 1):
        row_idx = i + 1
        row_cells = table.rows[row_idx].cells

        left_idx = i
        if left_idx < num_students:
            class1, name1 = students[left_idx]
            row_cells[0].text = str(left_idx + 1)
            row_cells[1].text = class1
            row_cells[2].text = name1

        right_idx = i + (num_rows - 1)
        if right_idx < num_students:
            class2, name2 = students[right_idx]
            row_cells[3].text = str(right_idx + 1)
            row_cells[4].text = class2
            row_cells[5].text = name2

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '等线'
                    run.font.size = Pt(11)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
    # 添加统计部分
    doc.add_paragraph("经整合：")
    class_counts = {}
    for class_name, _ in students:  # 已经按班级排序
        if class_name not in class_counts:
            class_counts[class_name] = 0
        class_counts[class_name] += 1

    num_classes = len(class_counts)
    stat_rows = (num_classes + 1) // 2 + 2
    stat_table = doc.add_table(rows=stat_rows, cols=6)
    stat_table.style = 'Table Grid'

    stat_headers = ["班级", "请假总人数", "备注", "班级", "请假总人数", "备注"]
    stat_hdr_cells = stat_table.rows[0].cells
    for i, header in enumerate(stat_headers):
        stat_hdr_cells[i].text = header
        for paragraph in stat_hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = '等线'
                run.font.size = Pt(11)
                run.bold = True
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
        set_cell_shading(stat_hdr_cells[i], "D9D9D9")

    class_list = list(class_counts.items())
    for i in range(0, len(class_list), 2):
        row_idx = i // 2 + 1
        row_cells = stat_table.rows[row_idx].cells

        class1, count1 = class_list[i]
        row_cells[0].text = class1
        row_cells[1].text = str(count1)

        if i + 1 < len(class_list):
            class2, count2 = class_list[i + 1]
            row_cells[3].text = class2
            row_cells[4].text = str(count2)

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '等线'
                    run.font.size = Pt(11)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

    total_row = stat_table.rows[stat_rows - 1].cells
    total_row[3].text = "总计"
    total_row[4].text = str(len(students))

    for i in [0, 1, 3, 4]:
        for paragraph in total_row[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = '等线'
                run.font.size = Pt(11)
                run.bold = True if i in [0, 3] else False
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
    set_cell_shading(total_row[3], "D9D9D9")
    doc.add_paragraph()
    signature = doc.add_paragraph("计算机与信息技术学院")
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in signature.runs:
        run.font.name = '等线'
        run.font.size = Pt(11)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

    date = doc.add_paragraph(f"{year}年{month}月{day}日")
    date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in date.runs:
        run.font.name = '等线'
        run.font.size = Pt(11)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

    set_cell_shading(total_row[3], "D9D9D9")
    doc.add_paragraph()
    signature = doc.add_paragraph("负责人签名：__________________")
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in signature.runs:
        run.font.name = '等线'
        run.font.size = Pt(11)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

    set_cell_shading(total_row[3], "D9D9D9")
    doc.add_paragraph()
    signature = doc.add_paragraph("指导老师签名：__________________")
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in signature.runs:
        run.font.name = '等线'
        run.font.size = Pt(11)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    file_name = f"{year}年{month}月{day}日_{cause}假单.docx"
    file_path = os.path.join(desktop_path, file_name)
    doc.save(file_path)
    print(f"请假单已生成：{cause}假单_{year}.{month}.{day}.docx")
    os.startfile(file_path)

if __name__ == "__main__":
    print("gitee：https://gitee.com/z_ky/Fake-orders.git")
    print("用于部门假单生成")
    print("数据来源：微信接龙复制/共享文档复制")
    print("注：若有错误请自行修改，上次更新时间：2025.11.13")
    input_time = input("请输入请假时间（例如：2025.4.27）：")
    cause = input("计信学院因xxx工作需要，以下同学需请假。(例：DH部)")
    input_time = input_time.split(".")
    year, month, day = int(input_time[0]), int(input_time[1]), int(input_time[2])
    create_leave_form()