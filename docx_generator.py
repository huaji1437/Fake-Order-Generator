import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Tuple, Any, LiteralString


class DocumentGenerator:
    """生成请假单文档的类"""
    config: Any
    def __init__(self, config_reader):
        self.config = config_reader

    def set_cell_shading(self, cell, shade: str):
        """设置单元格底纹颜色"""
        tcPr = cell._tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), shade)
        tcPr.append(shading)

    def apply_font_settings(self, run, font_type: str = "normal"):
        """应用字体设置"""
        font_settings = self.config.get("font_settings", {})
        font_name = font_settings.get(f"{font_type}_font", "等线")
        font_size = font_settings.get("font_size", {}).get(font_type, 11)

        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def create_leave_form(self, students: List[Tuple[str, str]], year: int, month: int, day: int,
                          cause: str, leave_type: str = "evening"):
        """创建请假单文档"""
        doc = Document()

        # 设置默认样式
        style = doc.styles['Normal']
        font = style.font
        normal_font = self.config.get("font_settings.normal_font", "等线")
        normal_size = self.config.get("font_settings.font_size.normal", 11)
        font.name = normal_font
        font.size = Pt(normal_size)
        font.element.rPr.rFonts.set(qn('w:eastAsia'), normal_font)

        # 添加标题
        self._add_title(doc, year, month, day, leave_type)

        # 添加正文内容
        self._add_content(doc, cause, students)

        # 添加统计表格
        self._add_statistics_table(doc, students)

        # 添加签名部分
        self._add_signature(doc, year, month, day)

        # 保存文件
        return self._save_document(doc, year, month, day, cause)

    def _add_title(self, doc: Document, year: int, month: int, day: int, leave_type: str):
        """添加标题部分"""
        college_name = self.config.get("college_name", "？？？？")
        title_format = self.config.get("title_format", "%s")

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(title_format % college_name)
        title_run.bold = True
        self.apply_font_settings(title_run, "title")

        # 时间部分
        time_para = doc.add_paragraph()
        time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        leave_types = self.config.get("leave_types", {})
        leave_text = leave_types.get(leave_type, "晚自习")
        time_text: str
        if leave_type == "morning":
            time_text = f"☑早自习 □晚自习 请假时间：{year}年{month}月{day}日"
        elif leave_type == "evening":
            time_text = f"□早自习 ☑晚自习 请假时间：{year}年{month}月{day}日"
        else:
            time_text = f"□早自习 ☑晚自习? 请假时间：{year}年{month}月{day}日"

        time_run = time_para.add_run(time_text)
        self.apply_font_settings(time_run, "small")

    def _add_content(self, doc: Document, cause: str, students: List[Tuple[str, str]]):
        """添加正文内容"""
        # 问候语
        greeting = doc.add_paragraph()
        greeting_run = greeting.add_run("各班级：")
        self.apply_font_settings(greeting_run, "content")

        # 请假原因
        reason = doc.add_paragraph()
        reason.paragraph_format.first_line_indent = Inches(0.3)
        reason_run = reason.add_run(f"因{cause}工作需要，以下同学需请假。")
        self.apply_font_settings(reason_run, "content")

        # 学生表格
        self._add_students_table(doc, students)

    def _add_students_table(self, doc: Document, students: List[Tuple[str, str]]):
        """添加学生信息表格"""
        students.sort(key=lambda x: (x[0][:2], x[0][2:]))
        num_students = len(students)
        num_rows = (num_students + 1) // 2 + 1

        table = doc.add_table(rows=num_rows, cols=6)
        table.style = 'Table Grid'

        headers = ["序号", "班级", "姓名", "序号", "班级", "姓名"]
        hdr_cells = table.rows[0].cells
        shading_color = self.config.get("table_settings.header_shading", "D9D9D9")

        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    self.apply_font_settings(run)
            self.set_cell_shading(hdr_cells[i], shading_color)

        # 填充学生数据
        for i in range(num_rows - 1):
            row_idx = i + 1
            row_cells = table.rows[row_idx].cells

            left_idx = i
            if left_idx < num_students:
                class1, name1 = students[left_idx]
                row_cells[0].text = str(left_idx + 1)
                row_cells[1].text = class1
                row_cells[2].text = name1

            right_idx = i + (num_rows - 1)
            if right_idx < num_students:
                class2, name2 = students[right_idx]
                row_cells[3].text = str(right_idx + 1)
                row_cells[4].text = class2
                row_cells[5].text = name2

            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        self.apply_font_settings(run)

    def _add_statistics_table(self, doc: Document, students: List[Tuple[str, str]]):
        """添加统计表格"""
        doc.add_paragraph("经整合：")

        # 统计各班人数
        class_counts = {}
        for class_name, _ in students:
            if class_name not in class_counts:
                class_counts[class_name] = 0
            class_counts[class_name] += 1

        num_classes = len(class_counts)
        stat_rows = (num_classes + 1) // 2 + 2
        stat_table = doc.add_table(rows=stat_rows, cols=6)
        stat_table.style = 'Table Grid'

        stat_headers = ["班级", "请假总人数", "备注", "班级", "请假总人数", "备注"]
        stat_hdr_cells = stat_table.rows[0].cells
        shading_color = self.config.get("table_settings.header_shading", "D9D9D9")

        for i, header in enumerate(stat_headers):
            stat_hdr_cells[i].text = header
            for paragraph in stat_hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    self.apply_font_settings(run)
            self.set_cell_shading(stat_hdr_cells[i], shading_color)

        # 填充统计数据
        class_list = list(class_counts.items())
        for i in range(0, len(class_list), 2):
            row_idx = i // 2 + 1
            row_cells = stat_table.rows[row_idx].cells

            class1, count1 = class_list[i]
            row_cells[0].text = class1
            row_cells[1].text = str(count1)

            if i + 1 < len(class_list):
                class2, count2 = class_list[i + 1]
                row_cells[3].text = class2
                row_cells[4].text = str(count2)

            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        self.apply_font_settings(run)

        # 总计行
        total_row = stat_table.rows[stat_rows - 1].cells
        total_row[3].text = "总计"
        total_row[4].text = str(len(students))

        for i in [0, 1, 3, 4]:
            for paragraph in total_row[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True if i in [0, 3] else False
                    self.apply_font_settings(run)
        self.set_cell_shading(total_row[3], shading_color)

    def _add_signature(self, doc: Document, year: int, month: int, day: int):
        """添加签名部分"""
        college_name = self.config.get("college_name", "？？？？")

        doc.add_paragraph()
        signature = doc.add_paragraph(college_name)
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in signature.runs:
            self.apply_font_settings(run)

        date = doc.add_paragraph(f"{year}年{month}月{day}日")
        date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in date.runs:
            self.apply_font_settings(run)

        doc.add_paragraph()
        responsible = doc.add_paragraph("负责人签名：__________________")
        responsible.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in responsible.runs:
            self.apply_font_settings(run)

        doc.add_paragraph()
        teacher = doc.add_paragraph("指导老师签名：__________________")
        teacher.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in teacher.runs:
            self.apply_font_settings(run)

    def _save_document(self, doc: Document, year: int, month: int, day: int, cause: str) -> LiteralString | str | bytes:
        """保存文档"""
        output_settings = self.config.get("output_settings", {})
        save_path = output_settings.get("save_path", "desktop")
        file_name_format = output_settings.get("file_name_format", "{year}年{month}月{day}日_{cause}假单.docx")

        if save_path == "desktop":
            save_path = os.path.join(os.path.expanduser("~"), "Desktop")

        file_name = file_name_format.format(year=year, month=month, day=day, cause=cause)
        file_path = os.path.join(save_path, file_name)

        doc.save(file_path)
        return file_path